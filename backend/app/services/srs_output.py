"""Validate srs.json, render md/docx/pdf, and write an immutable version folder atomically."""
from __future__ import annotations

import json
import logging
import shutil
from datetime import datetime, timezone
from pathlib import Path

from jsonschema import ValidationError, validate
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.agent.prompt import load_schema
from app.config import settings
from app.db.models import Project, SrsVersion

logger = logging.getLogger("ba-agent.srs_output")


_VALID_PRIORITIES = {"must", "should", "could", "wont"}
_PRIORITY_ALIASES = {
    "must have": "must", "must-have": "must", "high": "must", "critical": "must", "m": "must",
    "should have": "should", "should-have": "should", "medium": "should", "med": "should",
    "important": "should", "s": "should",
    "could have": "could", "could-have": "could", "low": "could", "nice to have": "could",
    "nice-to-have": "could", "optional": "could", "c": "could",
    "won't": "wont", "won't have": "wont", "wont have": "wont", "will not have": "wont",
    "out of scope": "wont", "w": "wont",
}
_VALID_REQ_TYPES = {"functional", "non_functional", "constraint", "assumption"}
_TYPE_ALIASES = {
    "non functional": "non_functional", "nonfunctional": "non_functional", "nfr": "non_functional",
    "fr": "functional", "func": "functional", "functional requirement": "functional",
    "constraints": "constraint", "assumptions": "assumption",
}


def _coerce_priority(v) -> str:
    s = str(v or "").strip().lower()
    if s in _VALID_PRIORITIES:
        return s
    return _PRIORITY_ALIASES.get(s, "should")


def _coerce_req_type(v) -> str:
    s = str(v or "").strip().lower()
    if s in _VALID_REQ_TYPES:
        return s
    normalized = s.replace("-", " ").replace("_", " ").strip()
    return _TYPE_ALIASES.get(s, _TYPE_ALIASES.get(normalized, "functional"))


def normalize_srs(data: dict) -> dict:
    """Fill sensible defaults and map common LLM variants so a single missing/loose field
    (e.g. a requirement without `priority`) doesn't throw away an otherwise-good run. The schema
    stays strict for genuinely broken output; this only heals the routine omissions."""
    if not isinstance(data, dict):
        return data
    data.setdefault("schema_version", "1.0")

    proj = data.get("project")
    if not isinstance(proj, dict):
        data["project"] = {"name": ""}
    else:
        proj.setdefault("name", "")

    reqs = data.get("requirements")
    reqs = reqs if isinstance(reqs, list) else []
    fixed: list[dict] = []
    for i, r in enumerate(reqs, start=1):
        if not isinstance(r, dict):
            continue
        r.setdefault("id", f"FR-{i:03d}")
        r["type"] = _coerce_req_type(r.get("type"))
        title = r.get("title") or r.get("name") or ""
        r["title"] = str(title).strip() or f"Requirement {r['id']}"
        desc = r.get("description")
        if desc is None:
            desc = r.get("summary") or ""
        r["description"] = str(desc)
        r["priority"] = _coerce_priority(r.get("priority"))
        fixed.append(r)
    data["requirements"] = fixed
    return data


def validate_srs(data: dict) -> None:
    try:
        validate(instance=data, schema=load_schema())
    except ValidationError as e:
        raise ValueError(f"SRS JSON failed validation: {e.message}") from e


def summarize(data: dict) -> str:
    reqs = data.get("requirements", []) or []
    oq = data.get("open_questions", []) or []
    nfr = data.get("non_functional", []) or []
    return f"{len(reqs)} requirements, {len(nfr)} non-functional, {len(oq)} open questions."


def render_markdown(data: dict) -> str:
    lines: list[str] = []
    project = data.get("project", {})
    lines.append(f"# Software Requirements Specification — {project.get('name', '')}")
    meta = data.get("meta", {})
    if meta:
        modules = ", ".join(meta.get("modules", []) or [])
        lines.append("")
        lines.append(f"**Domain:** {meta.get('domain', 'Odoo')}  ")
        if meta.get("odoo_version"):
            lines.append(f"**Odoo version:** {meta.get('odoo_version')}  ")
        if modules:
            lines.append(f"**Modules:** {modules}  ")
    if data.get("overview"):
        lines += ["", "## 1. Overview", "", data["overview"]]

    actors = data.get("actors", []) or []
    if actors:
        lines += ["", "## 2. Actors"]
        for a in actors:
            lines.append(f"- **{a.get('name','')}** — {a.get('role','')}")

    reqs = data.get("requirements", []) or []
    if reqs:
        lines += ["", "## 3. Requirements"]
        for r in reqs:
            lines += ["", f"### {r.get('id','')} — {r.get('title','')}"]
            lines.append(f"*Type:* {r.get('type','')} · *Priority:* {r.get('priority','')}"
                         + (f" · *Odoo module:* {r.get('odoo_module')}" if r.get("odoo_module") else ""))
            lines += ["", r.get("description", "")]
            ac = r.get("acceptance_criteria", []) or []
            if ac:
                lines.append("\n**Acceptance criteria:**")
                lines += [f"- {c}" for c in ac]
            refs = r.get("source_refs", []) or []
            if refs:
                lines.append("\n**Sources:**")
                lines += [f"- {ref.get('document','')}: \"{ref.get('quote','')}\"" for ref in refs]

    nfr = data.get("non_functional", []) or []
    if nfr:
        lines += ["", "## 4. Non-Functional Requirements"]
        for n in nfr:
            metric = f" (metric: {n.get('metric')})" if n.get("metric") else ""
            lines.append(f"- **{n.get('id','')}** [{n.get('category','')}] {n.get('description','')}{metric}")

    glossary = data.get("glossary", []) or []
    if glossary:
        lines += ["", "## 5. Glossary"]
        for g in glossary:
            lines.append(f"- **{g.get('term','')}** — {g.get('definition','')}")

    oq = data.get("open_questions", []) or []
    if oq:
        lines += ["", "## 6. Open Questions"]
        for q in oq:
            flag = " **(blocking)**" if q.get("blocking") else ""
            lines.append(f"- {q.get('id','')}: {q.get('question','')}{flag}")

    return "\n".join(lines) + "\n"


def render_docx(data: dict, md: str, out: Path) -> None:
    from docx import Document as Docx

    doc = Docx()
    for raw in md.splitlines():
        line = raw.rstrip()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(str(out))


def render_pdf(md: str, out: Path) -> None:
    try:
        import markdown as md_lib
        from weasyprint import HTML

        html = md_lib.markdown(md, extensions=["tables"])
        styled = f"""<html><head><meta charset='utf-8'><style>
        body{{font-family:Segoe UI,Arial,sans-serif;font-size:12px;line-height:1.5;margin:32px;}}
        h1{{font-size:22px;}} h2{{font-size:17px;margin-top:18px;}} h3{{font-size:14px;}}
        table{{border-collapse:collapse;}} td,th{{border:1px solid #ccc;padding:4px 8px;}}
        </style></head><body>{html}</body></html>"""
        HTML(string=styled).write_pdf(str(out))
        return
    except Exception as e:  # fall back to reportlab (pure-Python)
        logger.warning("WeasyPrint unavailable (%s); using reportlab fallback.", e)

    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    story = []
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        style = styles["BodyText"]
        text = line
        if line.startswith("# "):
            style, text = styles["Title"], line[2:]
        elif line.startswith("## "):
            style, text = styles["Heading2"], line[3:]
        elif line.startswith("### "):
            style, text = styles["Heading3"], line[4:]
        story.append(Paragraph(text.replace("&", "&amp;").replace("<", "&lt;"), style))
    SimpleDocTemplate(str(out), pagesize=A4).build(story)


def write_version(
    db: Session,
    *,
    project: Project,
    data: dict,
    model_id: str,
    template_id: str | None,
    source_docs_hash: str,
    generated_by: str,
    job_id: str,
) -> SrsVersion:
    normalize_srs(data)
    validate_srs(data)

    next_no = (db.scalar(
        select(func.max(SrsVersion.version_no)).where(SrsVersion.project_id == project.id)
    ) or 0) + 1

    # enrich source metadata
    data.setdefault("project", {})
    data["project"]["name"] = project.name
    data["project"]["srs_version"] = next_no
    data["project"]["generated_at"] = datetime.now(timezone.utc).isoformat()
    data.setdefault("source", {})
    data["source"]["model"] = model_id

    md = render_markdown(data)

    final_dir = settings.ba_output_dir / project.slug / f"v{next_no}"
    final_dir.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = final_dir.parent / f".v{next_no}.tmp"
    if tmp_path.exists():
        shutil.rmtree(tmp_path, ignore_errors=True)
    tmp_path.mkdir(parents=True)
    try:
        (tmp_path / "srs.json").write_text(json.dumps(data, indent=2), encoding="utf-8")
        (tmp_path / "srs.md").write_text(md, encoding="utf-8")
        render_docx(data, md, tmp_path / "srs.docx")
        render_pdf(md, tmp_path / "srs.pdf")

        if final_dir.exists():
            shutil.rmtree(final_dir, ignore_errors=True)
        tmp_path.rename(final_dir)  # atomic on same filesystem
    finally:
        shutil.rmtree(tmp_path, ignore_errors=True)

    version = SrsVersion(
        project_id=project.id,
        version_no=next_no,
        md_path=str(final_dir / "srs.md"),
        json_path=str(final_dir / "srs.json"),
        docx_path=str(final_dir / "srs.docx"),
        pdf_path=str(final_dir / "srs.pdf"),
        template_id=template_id,
        model_id=model_id,
        source_docs_hash=source_docs_hash,
        generated_by=generated_by,
        job_id=job_id,
        host_sync_status="not_sent",
    )
    db.add(version)
    db.flush()
    project.current_srs_version_id = version.id
    project.srs_status = "generated"
    project.last_generated_at = datetime.now(timezone.utc)
    db.commit()
    return version
